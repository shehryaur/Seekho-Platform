"""
seekho_v3.py  --  Seekho Platform
Pakistan's Hyper-Local AI Curriculum Engine

"""

import io
import json
import os
import re
import urllib.parse

import streamlit as st
from datetime import datetime
from google import genai

import database as db
import ui_style

# ── PAGE CONFIG ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Seekho Platform",
    page_icon=ui_style.LOGO_URL,
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={"About": "Seekho Platform -- PCTB-aligned AI curriculum for Pakistani teachers."},
)

# ── GLOBAL STYLES ─────────────────────────────────────────────────────────────
ui_style.inject_all()

# ── SESSION STATE ─────────────────────────────────────────────────────────────
_DEFAULTS = {
    # core
    "splash_done":        False,
    "lesson_json":        None,     # Feature 1: parsed dict from Gemini JSON response
    "lesson_content":     "",       # raw text fallback (for edit tab and copy)
    "last_params":        {},
    "history":            [],
    "wa_messages":        "",
    "just_generated":     False,
    "last_school":        "",
    "last_district":      "",
    # new
    "feedback_given":     False,    # Feature 9: prevent double submission
    "waitlist_submitted": False,    # Feature 8: hide form after submit
}
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

# ── SPLASH ────────────────────────────────────────────────────────────────────
ui_style.maybe_show_splash()

# ── GEMINI CLIENT ─────────────────────────────────────────────────────────────
@st.cache_resource
def _gemini():
    key = ""
    try:
        key = st.secrets["GEMINI_API_KEY"]
    except Exception:
        key = os.environ.get("GEMINI_API_KEY", "")
    if not key:
        st.error("GEMINI_API_KEY missing. Add it to .streamlit/secrets.toml and restart.")
        st.stop()
    return genai.Client(api_key=key)

gemini = _gemini()
MODEL  = "gemini-flash-latest"

# ── CLASS CONFIGS ─────────────────────────────────────────────────────────────
CLASS_CFG = {
    1:  {"vocab": "3-5 word sentences. Monosyllabic words. Heavy repetition.",
         "q":    "True/False or circle-the-picture. Max 3 questions.",
         "bl":   "REMEMBER only.",
         "len":  "60-90 words. One concept.",
         "n":    "Teacher reads aloud. Writing = tracing."},
    2:  {"vocab": "2-syllable words. 5-7 word sentences.",
         "q":    "Match column or 3-option MCQ. Max 4 questions.",
         "bl":   "REMEMBER + basic UNDERSTAND.",
         "len":  "100-140 words.",
         "n":    "Anchor to objects the child can hold."},
    3:  {"vocab": "Everyday words. Max 8 words/sentence. Repeat key terms 3 times.",
         "q":    "3-option MCQ + 1 fill-in-blank. 5 questions.",
         "bl":   "REMEMBER + UNDERSTAND. Simple categorisation.",
         "len":  "150-200 words.",
         "n":    "Group oral activities work better than individual writing."},
    4:  {"vocab": "Define any term inline. Max 10 words/sentence.",
         "q":    "3-option MCQ + match column. 5-6 questions.",
         "bl":   "REMEMBER, UNDERSTAND, simple APPLY.",
         "len":  "200-260 words.",
         "n":    "Drawing activities improve retention for this age."},
    5:  {"vocab": "Technical terms with bracket definition. Max 12 words.",
         "q":    "4-option MCQ + fill-blank + 1 short answer. 6 questions.",
         "bl":   "REMEMBER, UNDERSTAND, APPLY.",
         "len":  "250-320 words.",
         "n":    "PCTB board pressure begins here. Past-paper vocabulary."},
    6:  {"vocab": "Technical terms with parenthetical definition. Max 12 words.",
         "q":    "4-option MCQ + short answer + diagram label. 7 questions.",
         "bl":   "REMEMBER, UNDERSTAND, APPLY. One application question.",
         "len":  "300-380 words.",
         "n":    "BISE alignment critical. Match PCTB chapter references."},
    7:  {"vocab": "Board-level vocabulary. Sentences up to 14 words.",
         "q":    "MCQ (4-option) + SQs (2-3 marks) + 1 reasoning question. 8 questions.",
         "bl":   "REMEMBER through ANALYZE. At least one analytical question.",
         "len":  "350-430 words.",
         "n":    "Pair work and written discussion appropriate."},
    8:  {"vocab": "Full PCTB vocabulary. Board exam sentence structure.",
         "q":    "MCQ + SQ + LQ outline. 8-10 questions. Board format.",
         "bl":   "All levels through ANALYZE.",
         "len":  "400-500 words.",
         "n":    "Matric preparation. Formal, exam-aligned tone."},
    9:  {"vocab": "Complete PCTB Matric vocabulary. Precise definitions.",
         "q":    "MCQ (1-mark) + SQ (2-3 mark) + LQ (5-mark). 10 questions.",
         "bl":   "All Bloom's levels. Analysis + evaluation required.",
         "len":  "450-550 words. Board exam format.",
         "n":    "Reference PCTB chapter explicitly. Past-paper language."},
    10: {"vocab": "Board exam vocabulary. Complex sentences acceptable.",
         "q":    "Full board format: MCQ + SQ + LQ. 10-12 questions.",
         "bl":   "All levels. CREATE + EVALUATE dominant.",
         "len":  "500-620 words. BISE standard.",
         "n":    "Board year. Every element must be exam-relevant."},
    11: {"vocab": "University-entrance vocabulary. FSc/FA technical language.",
         "q":    "HSSC SQs + LQs. Essay-type possible. 8-12 questions.",
         "bl":   "EVALUATE + CREATE dominant. Cross-topic synthesis.",
         "len":  "600+ words. HSSC standard.",
         "n":    "MDCAT/ECAT relevance where applicable."},
    12: {"vocab": "University-entrance vocabulary. Critical analysis language.",
         "q":    "Board SQs, LQs, numericals. 10-14 questions.",
         "bl":   "EVALUATE, CREATE, synthesis.",
         "len":  "650+ words. HSSC final year.",
         "n":    "Entry test preparation is the primary driver."},
}


# ── FEATURE 1 + 4 + 5 + 6: JSON PROMPT BUILDER ───────────────────────────────
def build_prompt(
    school, district, class_num, subject, chapter, topic,
    language, profile, extra, topics_list,
    urdu_translate: bool = False,   # Feature 4
    low_resource: bool = False,     # Feature 5
) -> str:

    d     = db.get_district(district)
    cc    = CLASS_CFG.get(class_num, CLASS_CFG[6])
    today = datetime.now().strftime("%d %b %Y")

    # Language block
    lang_block = {
        "Pure Urdu (Script)": (
            "LANGUAGE: Write ALL sections in pure Urdu Nastaliq script. "
            "Only exception: scientific symbols (CO2, H2O, km, PKR)."
        ),
        "Roman Urdu": (
            "LANGUAGE: Write ALL sections in Roman Urdu (Urdu words in English letters), "
            "exactly as Pakistani students type on WhatsApp. NOT formal English. NOT Urdu script."
        ),
    }.get(language, (
        "LANGUAGE: Clear Pakistani English. Include Urdu or Punjabi terms in brackets where natural."
    ))

    # Profile block
    profile_block = {
        "Weak Class (Below Average)": (
            "CLASS PROFILE: Reduce vocabulary one full grade below stated class. "
            "No abstract concepts. Write for oral delivery. Repeat key term 3 times."
        ),
        "Strong Class (Above Average)": (
            "CLASS PROFILE: Add one [CHALLENGE] question above class level. "
            "Add a Did You Know fact beyond the PCTB chapter. Invite hypothesis formation."
        ),
    }.get(profile, "CLASS PROFILE: Standard complexity for Class {}.".format(class_num))

    # Feature 4: Urdu translation injection
    urdu_block = (
        "\nURDU TRANSLATION ACTIVE: The student_handbook, class_activity, and quiz.questions "
        "MUST be written in simple, readable Urdu (Nastaliq script). "
        "The teacher_guide stays in English. Do not mix scripts within a section.\n"
    ) if urdu_translate else ""

    # Feature 5: Low-resource injection
    low_resource_block = (
        "\nLOW-RESOURCE MODE ACTIVE: The class_activity MUST use ONLY items available "
        "for under 10 PKR -- paper, pen, chalk, string, plastic bottle, leaves, dirt, or stones. "
        "Absolutely no lab equipment, projectors, computers, or purchased kits. "
        "Any activity requiring such items is REJECTED.\n"
    ) if low_resource else ""

    topics_str = (
        "\n".join(f"  - {t}" for t in topics_list)
        if topics_list else "  - All topics in this chapter"
    )

    # Feature 1 + 6: JSON output schema with 40-minute constraint
    return f"""
You are SEEKHO ENGINE -- a specialist Pakistani curriculum expert, 20 years in PCTB-aligned schools.

PARAMETERS
School: {school} | District: {district} | Board: {d.get("board", "BISE")}
Class: {class_num} | Subject: {subject} | Chapter: {chapter}
Focus: {topic if topic not in ("Full Chapter Overview", "") else "Complete chapter"} | Date: {today}
Language: {language} | Profile: {profile}

LOCAL CONTEXT FOR {district.upper()}
Economy: {d.get("economy", "")} | Landmarks: {d.get("landmarks", "")}
Transport: {d.get("transport", "")} | Food: {d.get("food", "")}
Occupations: {d.get("occupations", "")} | Nature: {d.get("nature", "")}
Local Names: {d.get("local_names", "")} | School type: {d.get("school_type", "")}

LOCALISATION RULE: Every example, analogy, and word problem MUST use a specific name, place,
or item from the LOCAL CONTEXT above. Test: "Would a student from {district} immediately
recognise this from their daily life?" If no -- replace it.

CHAPTER TOPICS: {topics_str}
PCTB REF: PCTB {subject} Class {class_num}, Chapter: {chapter}

CLASS {class_num} RULES
Vocabulary: {cc["vocab"]}
Questions: {cc["q"]}
Bloom level: {cc["bl"]}
Length target: {cc["len"]}
Note: {cc["n"]}

{lang_block}
{profile_block}
{urdu_block}
{low_resource_block}
EXTRA INSTRUCTIONS: {extra.strip() if extra.strip() else "None."}

CONTENT QUALITY RULES (non-negotiable):
- NO paragraph walls anywhere. Use tables, numbered lists, bullet lists only.
- CONCEPT LADDER in student_handbook: Step 1 local analogy (zero jargon), Step 2 bridge, Step 3 PCTB definition.
- QUICK THINK: 2 in-class questions requiring genuine thinking, not copying.
- MCQ INTEGRITY: exactly one correct answer. Three plausible wrong distractors.
- BLOOM LABELS on every quiz question: (K) (U) (A) (AN) (E).
- NO em dashes anywhere in any section.

FEATURE 6 -- MANDATORY 40-MINUTE STRUCTURE:
Both teacher_guide and class_activity MUST include an explicit minute-by-minute timeline.
Format: "Min 0-5: [action]", "Min 5-15: [action]", etc.
The blocks MUST sum to exactly 40 minutes. No block may exceed 15 minutes.

CRITICAL OUTPUT FORMAT -- READ CAREFULLY:
Return ONLY a valid JSON object. No markdown fences. No text before or after the JSON.
Use this EXACT schema:

{{
  "teacher_guide": "Complete teacher guide in Markdown. MUST contain: (a) Common Misconception with exact correction sentence. (b) Prerequisites as a bullet list. (c) 40-minute delivery plan as a Markdown table with columns Time / Activity / What To Say or Do -- blocks must sum to 40 min. (d) No-tech classroom activity with step-by-step instructions. (e) Wrap-Up Script -- exact words to say.",

  "student_handbook": "Complete student handout in Markdown. MUST contain: (a) The Big Question hook. (b) Concept Ladder Step 1 / Step 2 / Step 3 with local analogy from {district}. (c) Key Terms table: Term | Simple Meaning | Example From {district}. (d) How It Works numbered list. (e) Quick Think 2 in-class questions.",

  "class_activity": "A focused 40-minute in-class activity in Markdown. MUST contain: (a) Activity title and duration (40 min). (b) Explicit minute-by-minute timeline (Min 0-5, Min 5-15, etc. summing to 40 min). (c) Materials list -- low-resource if that flag is active. (d) Step-by-step instructions. (e) AI Home Activity block with exact prompt to type into ChatGPT/Gemini, expected output, student task, artifact to bring to class, and time needed.",

  "quiz": {{
    "questions": "Quiz questions ONLY -- NO answers. In Markdown. Part A: MCQs with 4 options and Bloom labels. Part B: Short questions (2 marks each) with Bloom labels. Part C: Real-life application (5 marks) grounded entirely in {district} context.",
    "answer_key": "Answers only. Start with: TEACHER ANSWER KEY -- DO NOT DISTRIBUTE. List every answer numbered clearly."
  }}
}}

Generate now. Begin with {{ and end with }}. Nothing else outside the JSON.
"""


# ── FEATURE 1: JSON PARSER ────────────────────────────────────────────────────
def parse_lesson_json(raw: str) -> dict:
    """
    Robust JSON parser with fallback.
    Strips markdown fences, attempts json.loads(), validates required keys.
    On failure wraps raw text so the UI always has something to display.
    """
    text = raw.strip()
    # Strip ```json ... ``` or ``` ... ``` fences
    text = re.sub(r"^```(?:json)?\s*\n?", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n?```\s*$",          "", text, flags=re.MULTILINE)
    text = text.strip()

    try:
        data = json.loads(text)
        required = {"teacher_guide", "student_handbook", "class_activity", "quiz"}
        if required.issubset(data.keys()):
            return data
    except (json.JSONDecodeError, ValueError, TypeError):
        pass

    # Graceful fallback -- never crash the UI
    return {
        "teacher_guide":    raw,
        "student_handbook": raw,
        "class_activity":   raw,
        "quiz": {
            "questions":   raw,
            "answer_key":  "(JSON parsing failed -- please regenerate the lesson.)",
        },
        "_fallback": True,
    }


# ── FEATURE 2: WHATSAPP SHARE TEXT BUILDER ────────────────────────────────────
def build_wa_share_text(lesson: dict, params: dict) -> str:
    """
    Compiles student_handbook + class_activity + quiz.questions ONLY.
    Excludes teacher_guide and answer_key as required.
    """
    school  = params.get("school_name", "Seekho Platform")
    subject = params.get("subject", "")
    chapter = params.get("chapter", "")
    cls     = params.get("class_num", "")

    header = (
        f"*{school}*\n"
        f"Class {cls} {subject} -- {chapter}\n"
        f"{datetime.now().strftime('%d %b %Y')}\n"
        f"{'='*38}\n\n"
    )
    sections = [
        lesson.get("student_handbook", ""),
        lesson.get("class_activity", ""),
        lesson.get("quiz", {}).get("questions", ""),
    ]
    body   = "\n\n---\n\n".join(s for s in sections if s.strip())
    footer = "\n\n---\n_Generated free by Seekho Platform_"
    return header + body + footer


# ── AI WHATSAPP MESSAGES ──────────────────────────────────────────────────────
def gen_whatsapp_messages(lesson: dict, params: dict) -> str:
    """AI-generated 3-audience WhatsApp messages (parent / student / teacher)."""
    p       = params
    snippet = lesson.get("student_handbook", "")[:900]
    prompt  = f"""
Format this lesson for WhatsApp. Generate 3 short messages separated by "---".

Lesson: Class {p.get('class_num','')} {p.get('subject','')} | {p.get('chapter','')} | {p.get('district','')}
School: {p.get('school_name','')}

MESSAGE A -- PARENT GROUP:
Open: "Assalam o Alaikum {p.get('school_name','')} parents!"
2 sentences what was taught. 1 question for parent to ask the child tonight. Under 130 words.

MESSAGE B -- STUDENT GROUP:
Casual elder-sibling tone. Key point in 2 fun sentences. 1 reply question. Under 90 words.

MESSAGE C -- TEACHER COLLEAGUES:
Professional. 1 teaching tip that works for this topic. Under 70 words.

Source content:
{snippet}
"""
    try:
        return gemini.models.generate_content(model=MODEL, contents=prompt).text
    except Exception as e:
        return f"Generation failed: {e}"


# ── FEATURE 3: DOCX FROM JSON ─────────────────────────────────────────────────
def _add_run(para, text: str) -> None:
    """Parse **bold** markers into docx runs."""
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            r.bold = True
        elif part:
            para.add_run(part)


def _md_to_docx_section(doc, md_text: str, GREEN, BLUE) -> None:
    """Render a Markdown string into an existing python-docx Document."""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    in_table = False
    t_hdrs: list = []
    t_rows: list = []

    def flush_table():
        nonlocal in_table, t_hdrs, t_rows
        if not t_hdrs:
            in_table = False
            return
        t = doc.add_table(rows=1 + len(t_rows), cols=len(t_hdrs))
        t.style = "Table Grid"
        for j, h in enumerate(t_hdrs):
            c = t.rows[0].cells[j]
            c.text = h.strip()
            for para in c.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        for ri, row in enumerate(t_rows, 1):
            for j, ct in enumerate(row):
                if j < len(t_hdrs):
                    t.rows[ri].cells[j].text = ct.strip()
        doc.add_paragraph()
        in_table = False
        t_hdrs.clear()
        t_rows.clear()

    for line in md_text.split("\n"):
        s = line.strip()

        # Table rows
        if s.startswith("|") and s.endswith("|"):
            parts = [p.strip() for p in s.strip("|").split("|")]
            if not in_table:
                in_table = True
                t_hdrs = parts
            elif re.match(r"^[\s\-|:]+$", s):
                pass  # separator row
            else:
                t_rows.append(parts)
            continue
        elif in_table:
            flush_table()

        if not s:
            doc.add_paragraph()
            continue
        if s.startswith("# "):
            h = doc.add_heading(s[2:], 1)
            if h.runs: h.runs[0].font.color.rgb = GREEN
        elif s.startswith("## "):
            h = doc.add_heading(s[3:], 2)
            if h.runs: h.runs[0].font.color.rgb = GREEN
        elif s.startswith("### "):
            h = doc.add_heading(s[4:], 3)
            if h.runs: h.runs[0].font.color.rgb = BLUE
        elif s.startswith(("* ", "- ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_run(p, s[2:])
        elif re.match(r"^\d+\.\s", s):
            p = doc.add_paragraph(style="List Number")
            _add_run(p, re.sub(r"^\d+\.\s*", "", s))
        elif s.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            r = p.add_run(s.lstrip("> "))
            r.italic = True
            r.font.size = Pt(11)
        elif s.startswith("---") or s.startswith("==="):
            doc.add_paragraph("_" * 54)
        elif s.startswith("*") and s.endswith("*") and not s.startswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(s.strip("*"))
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        else:
            p = doc.add_paragraph()
            _add_run(p, s)
            for r in p.runs:
                r.font.size = Pt(11)

    if in_table:
        flush_table()


def generate_docx(lesson: dict, params: dict) -> bytes:
    """Build a DOCX from the parsed JSON lesson dict. All sections included."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1.0)
        sec.left_margin = sec.right_margin = Inches(1.25)

    GREEN = RGBColor(0x1a, 0x7f, 0x4b)
    BLUE  = RGBColor(0x2a, 0x5f, 0xa0)
    RED   = RGBColor(0xc0, 0x39, 0x2b)

    # Cover block
    hp = doc.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(params.get("school_name", "Seekho Platform"))
    hr.bold = True
    hr.font.size = Pt(18)
    hr.font.color.rgb = GREEN

    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = mp.add_run(
        f"Class {params.get('class_num','')} | {params.get('subject','')} | "
        f"{params.get('chapter','')} | {params.get('district','')} | "
        f"{datetime.now().strftime('%d %b %Y')}"
    )
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # Render each section
    sections = [
        ("SECTION 1: TEACHER GUIDE",            lesson.get("teacher_guide", ""),                    GREEN),
        ("SECTION 2: STUDENT HANDBOOK",          lesson.get("student_handbook", ""),                 GREEN),
        ("SECTION 3: CLASS ACTIVITY (40 Min)",   lesson.get("class_activity", ""),                   GREEN),
        ("SECTION 4: QUIZ -- QUESTIONS",         lesson.get("quiz", {}).get("questions", ""),        GREEN),
        ("SECTION 4: QUIZ -- ANSWER KEY",        lesson.get("quiz", {}).get("answer_key", ""),       RED),
    ]

    for title, md_text, color in sections:
        if not md_text.strip():
            continue
        doc.add_paragraph("=" * 54)
        h = doc.add_heading(title, level=1)
        if h.runs:
            h.runs[0].font.color.rgb = color
        doc.add_paragraph()
        _md_to_docx_section(doc, md_text, GREEN, BLUE)
        doc.add_page_break()

    # Footer
    for sec in doc.sections:
        fp = sec.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(
            f"Seekho Platform | Let's make learning fun! | "
            f"{params.get('school_name', '')} | "
            f"Class {params.get('class_num', '')} {params.get('subject', '')} | "
            f"{datetime.now().strftime('%d %b %Y')}"
        )
        fr.font.size = Pt(8)
        fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── AI REFINE ─────────────────────────────────────────────────────────────────
def refine_section(text: str, instruction: str) -> str:
    return gemini.models.generate_content(model=MODEL, contents=f"""
Edit this Pakistani curriculum text.

INSTRUCTION: {instruction}

RULES:
- Apply ONLY the stated instruction. Change nothing else.
- Keep all Markdown formatting intact.
- No paragraph walls -- tables and lists only.
- No em dashes.
- Return the complete updated text with no preamble.

TEXT:
{text}
""").text


# ════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    ui_style.render_sidebar_logo()

    # Usage stats from analytics table (Feature 10)
    try:
        stats = db.get_analytics_summary()
        total = stats.get("total_lessons", 0)
        if total:
            st.markdown(
                f'<div style="padding:.4rem 0 .2rem;">'
                f'<span class="stat-pill">Lessons generated: {total:,}</span>'
                f'</div>',
                unsafe_allow_html=True,
            )
    except Exception:
        pass

    st.markdown("## Settings")

    class_profile = st.radio("Class Profile", [
        "Average Class",
        "Weak Class (Below Average)",
        "Strong Class (Above Average)",
    ])

    # Feature 4: Urdu Translation Toggle
    urdu_translate = st.checkbox(
        "Translate Student Materials to Simple Urdu",
        help=(
            "Student Handbook, Class Activity, and Quiz will be in simple Urdu. "
            "Teacher Guide stays in English."
        ),
    )

    # Feature 5: Low-Resource Activity
    low_resource = st.checkbox(
        "Use Only Low-Resource Materials",
        help=(
            "Forces all activities to use items under 10 PKR only: "
            "paper, string, chalk, plastic bottles, leaves, dirt, stones. "
            "No lab equipment."
        ),
    )

    extra = st.text_area(
        "Custom Instructions",
        placeholder=(
            "Use cricket analogies\n"
            "Add Ramadan context\n"
            "Include drawing activity\n"
            "Reference Waris Shah"
        ),
        height=100,
    )

    st.divider()

    # Session history
    st.markdown("## History")
    if st.session_state.history:
        for i, h in enumerate(reversed(st.session_state.history[-6:])):
            lbl = f"**{h.get('chapter', 'Lesson')}** (Cl.{h.get('class_num', '')})"
            with st.expander(lbl, expanded=False):
                st.caption(
                    f"{h.get('subject','')} | {h.get('district','')} | {h.get('time','')}"
                )
                if st.button("Reload", key=f"rl_{i}_{h.get('time', i)}"):
                    st.session_state.lesson_json    = h.get("lesson_json")
                    st.session_state.lesson_content = h.get("content_preview", "")
                    st.session_state.last_params    = h
                    st.session_state.feedback_given = False
                    st.rerun()
    else:
        st.caption("No lessons yet this session.")

    st.divider()

    # Feature 8: Waitlist with session_state hiding
    st.markdown("### Pro Waitlist")
    st.caption("Direct WhatsApp broadcast, prompt vault, school analytics -- coming soon.")

    if not st.session_state.waitlist_submitted:
        wl_ph = st.text_input(
            "WhatsApp Number", placeholder="03XX-XXXXXXX", label_visibility="collapsed"
        )
        wl_sc = st.text_input(
            "School Name", placeholder="School name", label_visibility="collapsed"
        )
        if st.button("Join Waitlist", use_container_width=True):
            if wl_ph.strip():
                try:
                    db.save_waitlist_entry(
                        wl_ph.strip(),
                        wl_sc.strip(),
                        st.session_state.last_params.get("district", ""),
                    )
                except Exception:
                    pass
                st.session_state.waitlist_submitted = True
                st.rerun()
            else:
                st.warning("Enter your WhatsApp number.")
    else:
        st.success("Added to waitlist!")


# ════════════════════════════════════════════════════════════════════════════
# MAIN HEADER
# ════════════════════════════════════════════════════════════════════════════
ui_style.render_main_header()
st.divider()


# ════════════════════════════════════════════════════════════════════════════
# EMPTY STATE
# ════════════════════════════════════════════════════════════════════════════
if not st.session_state.lesson_json:
    st.markdown("### What Seekho Platform creates for you")
    fc1, fc2, fc3, fc4 = st.columns(4)
    cards = [
        ("", "PCTB-Aligned Lesson",
         "Chapter-specific content mapped to your exact PCTB syllabus for Classes 1-12."),
        ("", "Hyper-Local Examples",
         "Every analogy uses real places, food, and occupations from your district."),
        ("", "4-Tab Output",
         "Teacher Guide, Student Handbook, 40-Min Activity, and Quiz -- each in its own clean tab."),
        ("", "Download + Share",
         "DOCX for the print shop and a direct WhatsApp share link for student materials."),
    ]
    for col, (icon, title, desc) in zip([fc1, fc2, fc3, fc4], cards):
        col.markdown(
            f'<div class="feature-card"><span class="fc-icon">{icon}</span>'
            f'<h4>{title}</h4><p>{desc}</p></div>',
            unsafe_allow_html=True,
        )
    st.markdown("")
    st.info("Tip: Press **Ctrl+Enter** after filling the form to generate instantly.")
    st.divider()


# ════════════════════════════════════════════════════════════════════════════
# INPUT FORM
# ════════════════════════════════════════════════════════════════════════════
ca, cb = st.columns(2)
with ca:
    school_name = st.text_input(
        "School Name",
        value=st.session_state.last_school,
        placeholder="e.g. Govt. Boys High School Fateh Jang",
    )
with cb:
    all_districts = db.get_district_names() + ["Other (type below)"]
    d_idx  = (
        all_districts.index(st.session_state.last_district)
        if st.session_state.last_district in all_districts else 0
    )
    d_pick = st.selectbox("District / City", all_districts, index=d_idx)
    district_name = (
        st.text_input("Custom district", placeholder="e.g. Chakwal")
        if d_pick == "Other (type below)" else d_pick
    )

c1, c2, c3 = st.columns(3)
with c1:
    class_num = st.selectbox("Class", range(1, 13), format_func=lambda x: f"Class {x}")
with c2:
    subjects = db.get_subjects(class_num)
    subject  = st.selectbox("Subject", subjects)
with c3:
    language = st.select_slider(
        "Language", options=["English", "Roman Urdu", "Pure Urdu (Script)"]
    )

c4, c5 = st.columns(2)
with c4:
    chapters       = db.get_chapters(class_num, subject)
    chapter        = st.selectbox("Chapter", chapters)
with c5:
    chapter_topics = db.get_topics(class_num, subject, chapter)

    # Feature 7: Custom Topic Input
    CUSTOM_SENTINEL = "Add Custom Topic..."
    topic_opts      = ["Full Chapter Overview"] + chapter_topics + [CUSTOM_SENTINEL]
    topic_select    = st.selectbox("Specific Topic (optional)", topic_opts)

    if topic_select == CUSTOM_SENTINEL:
        custom_topic = st.text_input(
            "Type your custom topic",
            placeholder="e.g. Kidney failure and dialysis in Pakistan",
        )
        topic = custom_topic.strip() or "Full Chapter Overview"
    else:
        topic = topic_select

# Active-flag display in preview bar
flags = []
if urdu_translate: flags.append("Urdu Mode")
if low_resource:   flags.append("Low-Resource")
flag_str = " | ".join(flags)

if school_name and district_name:
    st.markdown(
        f'<div class="meta-bar">Ready: <b>Class {class_num} {subject}</b>'
        f' | <b>{chapter}</b>'
        f'{(" | " + topic) if topic != "Full Chapter Overview" else ""}'
        f' | {district_name} | {language}'
        f'{(" | " + flag_str) if flag_str else ""}</div>',
        unsafe_allow_html=True,
    )

# Action buttons
btn_col, reset_col, _ = st.columns([1, 1, 6])
with btn_col:
    generate = st.button(
        "Generate", type="primary", use_container_width=True, help="Ctrl+Enter also works"
    )
with reset_col:
    if st.button("New Lesson", use_container_width=True, help="Clear and start fresh"):
        st.session_state.lesson_json    = None
        st.session_state.lesson_content = ""
        st.session_state.last_params    = {}
        st.session_state.wa_messages    = ""
        st.session_state.just_generated = False
        st.session_state.feedback_given = False
        st.rerun()


# ════════════════════════════════════════════════════════════════════════════
# VALIDATION
# ════════════════════════════════════════════════════════════════════════════
if generate:
    if not school_name.strip():
        st.warning("Please enter the school name.")
        generate = False
    elif not district_name or district_name.strip() == "Other (type below)":
        st.warning("Please select or type a district name.")
        generate = False
    elif topic_select == CUSTOM_SENTINEL and not custom_topic.strip():
        st.warning("Please type your custom topic, or choose from the dropdown.")
        generate = False


# ════════════════════════════════════════════════════════════════════════════
# GENERATION
# ════════════════════════════════════════════════════════════════════════════
if generate:
    prompt = build_prompt(
        school=school_name,
        district=district_name,
        class_num=class_num,
        subject=subject,
        chapter=chapter,
        topic=topic,
        language=language,
        profile=class_profile,
        extra=extra,
        topics_list=chapter_topics,
        urdu_translate=urdu_translate,
        low_resource=low_resource,
    )

    with st.spinner(f"Building lesson for {chapter} (Class {class_num}, {district_name})..."):
        try:
            raw_text = gemini.models.generate_content(model=MODEL, contents=prompt).text
            lesson   = parse_lesson_json(raw_text)

            if lesson.get("_fallback"):
                st.warning(
                    "The AI returned unstructured text instead of JSON. "
                    "Content is shown in all tabs. Try regenerating for better results."
                )

            params = {
                "school_name":    school_name,
                "district":       district_name,
                "class_num":      class_num,
                "subject":        subject,
                "chapter":        chapter,
                "topic":          topic,
                "language":       language,
                "class_profile":  class_profile,
                "urdu_translate": urdu_translate,
                "low_resource":   low_resource,
                "time":           datetime.now().strftime("%H:%M"),
                "content_preview": raw_text[:400],
                "lesson_json":    lesson,
            }

            st.session_state.lesson_json    = lesson
            st.session_state.lesson_content = raw_text
            st.session_state.last_params    = params
            st.session_state.last_school    = school_name
            st.session_state.last_district  = district_name
            st.session_state.wa_messages    = ""
            st.session_state.just_generated = True
            st.session_state.feedback_given = False

            # History (metadata only -- no full content blob to save memory)
            st.session_state.history.append(params)
            if len(st.session_state.history) > 10:
                st.session_state.history = st.session_state.history[-10:]

            # Feature 10: Usage Analytics (silent, never blocks UI)
            try:
                db.log_analytics(
                    district=district_name,
                    class_num=class_num,
                    subject=subject,
                    chapter=chapter,
                    language=language,
                    output_mode="Full Lesson Pack",
                )
            except Exception:
                pass

            # Existing: save full lesson to generated_lessons
            try:
                db.save_lesson(
                    school_name=school_name, district=district_name,
                    class_num=class_num, subject=subject, chapter=chapter,
                    topic=topic, language=language, output_mode="Full Lesson Pack",
                    class_profile=class_profile, content=raw_text,
                )
            except Exception:
                pass

            ui_style.play_success_sound()
            st.rerun()

        except Exception as e:
            st.error(f"Generation failed: {e}")


# ════════════════════════════════════════════════════════════════════════════
# OUTPUT
# ════════════════════════════════════════════════════════════════════════════
if st.session_state.lesson_json:
    lesson = st.session_state.lesson_json
    p      = st.session_state.last_params

    st.divider()
    st.markdown('<div class="output-wrap">', unsafe_allow_html=True)

    # ── Metadata strip + DOCX button ─────────────────────────────────────
    m1, m2, m3, m4, m5, d_col = st.columns([1.5, 1.5, 1.7, 1.5, 0.9, 1.5])
    m1.metric("Class",    f"Class {p.get('class_num','')}")
    m2.metric("Subject",  str(p.get("subject",""))[:18])
    m3.metric("Chapter",  str(p.get("chapter",""))[:20])
    m4.metric("District", str(p.get("district",""))[:16])

    total_words = sum(
        len(str(v).split())
        for v in [
            lesson.get("teacher_guide", ""),
            lesson.get("student_handbook", ""),
            lesson.get("class_activity", ""),
            str(lesson.get("quiz", "")),
        ]
    )
    m5.metric("Words", str(total_words))

    # Feature 3: DOCX download button inline (not a separate tab)
    with d_col:
        st.markdown("<br>", unsafe_allow_html=True)
        do_docx = st.button(
            "Download DOCX", use_container_width=True,
            help="Download full lesson as Word document for print shop"
        )

    if do_docx:
        with st.spinner("Building Word document..."):
            try:
                docx_bytes = generate_docx(lesson, p)
                fname = (
                    f"Seekho_{p.get('subject','Lesson').replace(' ','_')}_"
                    f"Cl{p.get('class_num','')}_"
                    f"{p.get('chapter','Ch').replace(' ','_')[:18]}.docx"
                )
                st.download_button(
                    "Click to download DOCX",
                    docx_bytes,
                    fname,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_docx",
                )
                st.toast("Word document ready!")
            except Exception as e:
                st.error(f"DOCX failed: {e}. Ensure python-docx is installed.")

    st.markdown("")

    # ── Feature 1: 6-Tab Output ──────────────────────────────────────────
    tab_tg, tab_sh, tab_ca, tab_qz, tab_wa, tab_rf = st.tabs([
        "📋 Teacher Guide",
        "📄 Student Handbook",
        "⚡ Class Activity",
        "📝 Quiz",
        "📱 WhatsApp",
        "🔧 AI Refine",
    ])

    # ── TAB 1: TEACHER GUIDE ─────────────────────────────────────────────
    with tab_tg:
        tg_text = lesson.get("teacher_guide", "")
        hcol, ccol = st.columns([6, 1])
        with ccol:
            ui_style.copy_button(tg_text, "copy_tg")
        st.info("Private -- do not distribute to students.")
        if p.get("language") == "Pure Urdu (Script)":
            st.markdown(f'<div class="urdu-rtl">{tg_text}</div>', unsafe_allow_html=True)
        else:
            st.markdown(tg_text)

    # ── TAB 2: STUDENT HANDBOOK ──────────────────────────────────────────
    with tab_sh:
        sh_text = lesson.get("student_handbook", "")
        hcol, ccol = st.columns([6, 1])
        with ccol:
            ui_style.copy_button(sh_text, "copy_sh")
        st.success("Print one copy per student. Press Ctrl+P to save as PDF.")
        is_urdu = p.get("language") == "Pure Urdu (Script)" or p.get("urdu_translate")
        if is_urdu:
            st.markdown(f'<div class="urdu-rtl">{sh_text}</div>', unsafe_allow_html=True)
        else:
            st.markdown(sh_text)

    # ── TAB 3: CLASS ACTIVITY ────────────────────────────────────────────
    with tab_ca:
        ca_text = lesson.get("class_activity", "")
        hcol, ccol = st.columns([6, 1])
        with ccol:
            ui_style.copy_button(ca_text, "copy_ca")
        is_urdu = p.get("language") == "Pure Urdu (Script)" or p.get("urdu_translate")
        if is_urdu:
            st.markdown(f'<div class="urdu-rtl">{ca_text}</div>', unsafe_allow_html=True)
        else:
            st.markdown(ca_text)

    # ── TAB 4: QUIZ ──────────────────────────────────────────────────────
    with tab_qz:
        quiz    = lesson.get("quiz", {})
        q_text  = quiz.get("questions", "")
        ak_text = quiz.get("answer_key", "")

        q_tab, ak_tab = st.tabs(["Questions (Student Copy)", "Answer Key (Teacher Only)"])

        with q_tab:
            hcol, ccol = st.columns([6, 1])
            with ccol:
                ui_style.copy_button(q_text, "copy_quiz")
            is_urdu = p.get("language") == "Pure Urdu (Script)" or p.get("urdu_translate")
            if is_urdu:
                st.markdown(f'<div class="urdu-rtl">{q_text}</div>', unsafe_allow_html=True)
            else:
                st.markdown(q_text)

        with ak_tab:
            st.warning("TEACHER ANSWER KEY -- Do not distribute to students.")
            st.markdown(ak_text)

    # ── TAB 5: WHATSAPP ──────────────────────────────────────────────────
    with tab_wa:
        st.markdown("### WhatsApp Distribution")

        # Feature 2: Direct wa.me share button (student content only)
        st.markdown("#### Share Student Materials Directly")
        st.caption(
            "Compiles Student Handbook + Class Activity + Quiz Questions. "
            "Teacher Guide and Answer Key are excluded."
        )
        wa_share_text = build_wa_share_text(lesson, p)
        wa_url        = f"https://wa.me/?text={urllib.parse.quote(wa_share_text)}"

        st.markdown(
            f'<a href="{wa_url}" target="_blank" rel="noopener noreferrer" style="'
            "display:inline-flex;align-items:center;gap:8px;"
            "background:#25D366;color:white;"
            "border-radius:10px;padding:10px 22px;"
            "font-size:.92rem;font-weight:700;"
            "text-decoration:none;"
            "font-family:'Plus Jakarta Sans',sans-serif;"
            "box-shadow:0 3px 14px rgba(37,211,102,.38);"
            '">'
            " Share to WhatsApp"
            "</a>",
            unsafe_allow_html=True,
        )

        st.divider()

        # AI-generated 3-audience messages
        st.markdown("#### AI-Generated Group Messages")
        st.markdown(
            "Three formatted messages for parent, student, and teacher groups."
        )

        if st.button("Generate Group Messages", type="primary", key="wa_btn"):
            with st.spinner("Formatting for WhatsApp..."):
                st.session_state.wa_messages = gen_whatsapp_messages(lesson, p)
                st.toast("Messages ready!")

        if st.session_state.wa_messages:
            st.markdown(st.session_state.wa_messages)
            ui_style.copy_button(st.session_state.wa_messages, "copy_wa")
            st.caption("Tap inside a message, select all, copy, paste into WhatsApp.")

        st.divider()
        st.markdown(
            "**Pro coming soon:** direct broadcast to contact lists, open-rate tracking, "
            "scheduled sends. Join the waitlist in the sidebar."
        )

    # ── TAB 6: AI REFINE ─────────────────────────────────────────────────
    with tab_rf:
        st.markdown("### AI Refinement Studio")

        dist = p.get("district", "the local area")

        # Section selector
        refine_target = st.selectbox(
            "Refine which section?",
            ["Student Handbook", "Class Activity", "Teacher Guide", "Quiz Questions"],
        )
        section_key_map = {
            "Teacher Guide":    "teacher_guide",
            "Student Handbook": "student_handbook",
            "Class Activity":   "class_activity",
            "Quiz Questions":   "quiz.questions",
        }
        s_key = section_key_map[refine_target]

        st.markdown("**Quick Actions:**")
        ra, rb, rc = st.columns(3)
        rd, re_, rf_ = st.columns(3)
        quick = None

        with ra:
            if st.button("More Local Examples", use_container_width=True):
                quick = (
                    f"Replace every generic analogy with something specific to {dist}, Pakistan. "
                    "Use real places, occupations, food. Name specific people and locations."
                )
        with rb:
            if st.button("Simplify Language", use_container_width=True):
                quick = (
                    "Reduce complexity by one grade level. Max 10 words per sentence. "
                    "Replace every technical word with a simpler alternative. Facts unchanged."
                )
        with rc:
            if st.button("Add Bloom Labels", use_container_width=True):
                quick = (
                    "Add Bloom labels to every question: "
                    "(K) Remember (U) Understand (A) Apply (AN) Analyze (E) Evaluate."
                )
        with rd:
            if st.button("Add Drawing Activity", use_container_width=True):
                quick = "Add a Draw and Label activity. Needs only pen and paper."
        with re_:
            if st.button("Double the MCQs", use_container_width=True):
                quick = (
                    "Double the MCQ count. Each new MCQ: one correct answer, "
                    "three plausible distractors. Span all Bloom levels. Local examples only."
                )
        with rf_:
            if st.button("Roman Urdu Version", use_container_width=True):
                quick = "Rewrite this section entirely in Roman Urdu (Urdu words in English letters)."

        st.markdown("")
        custom_inst = st.text_input(
            "Custom instruction",
            placeholder="e.g. Add Hadith reference | Add comparison table | Use farming analogies",
            label_visibility="collapsed",
        )

        ac, _ = st.columns([1, 4])
        with ac:
            if st.button("Apply", type="primary", use_container_width=True):
                instruction = quick or custom_inst
                if instruction:
                    # Get current text for the selected section
                    if s_key == "quiz.questions":
                        current_text = lesson.get("quiz", {}).get("questions", "")
                    else:
                        current_text = lesson.get(s_key, "")

                    with st.spinner("Applying..."):
                        try:
                            refined = refine_section(current_text, instruction)
                            # Update only the targeted section in session state
                            updated = dict(st.session_state.lesson_json)
                            if s_key == "quiz.questions":
                                updated["quiz"] = dict(updated.get("quiz", {}))
                                updated["quiz"]["questions"] = refined
                            else:
                                updated[s_key] = refined
                            st.session_state.lesson_json = updated
                            st.toast(f"{refine_target} updated!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Refinement failed: {e}")
                else:
                    st.warning("Select a quick action or type a custom instruction first.")

    # ── Feature 9: AI Feedback Loop ──────────────────────────────────────
    st.divider()
    st.markdown("**Was this lesson useful?**")

    if not st.session_state.feedback_given:
        fb1, fb2, _ = st.columns([1, 1, 6])
        with fb1:
            if st.button("👍 Good Lesson", use_container_width=True):
                try:
                    db.log_feedback(
                        rating="good",
                        district=p.get("district", ""),
                        class_num=p.get("class_num", 0),
                        subject=p.get("subject", ""),
                        chapter=p.get("chapter", ""),
                        topic=p.get("topic", ""),
                        language=p.get("language", ""),
                        lesson_json=json.dumps(lesson),
                    )
                except Exception:
                    pass
                st.session_state.feedback_given = True
                st.toast("Thanks! This helps us improve the engine.")
                st.rerun()
        with fb2:
            if st.button("👎 Bad Lesson", use_container_width=True):
                try:
                    db.log_feedback(
                        rating="bad",
                        district=p.get("district", ""),
                        class_num=p.get("class_num", 0),
                        subject=p.get("subject", ""),
                        chapter=p.get("chapter", ""),
                        topic=p.get("topic", ""),
                        language=p.get("language", ""),
                        lesson_json=json.dumps(lesson),
                    )
                except Exception:
                    pass
                st.session_state.feedback_given = True
                st.toast("Thanks! We'll use this to fix the engine.")
                st.rerun()
    else:
        st.caption("Thanks for your feedback!")

    st.markdown('</div>', unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════════════════════
# FOOTER
# ════════════════════════════════════════════════════════════════════════════
st.divider()
st.caption(
    "Seekho Platform | PCTB-Aligned | Free for all Pakistani teachers | "
    "AI-generated content -- review before use | Not affiliated with PCTB or any government body."
    "Email us: [shehryar.hassan@uni.minerva.edu](mailto:shehryar.hassan@uni.minerva.edu)"
)